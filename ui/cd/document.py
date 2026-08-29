"""
بناء وتوليد مستند "Change Devise" بصيغة Word، مطابق حرفياً (خط، حجم،
هوامش، مسافات، تنسيق الأرقام) لنموذج البوردرو الأصلي (Model Change
Devise.pdf) — القياسات مستخرجة رقمياً من ملف الـ PDF نفسه (نوع الخط،
حجمه، وموضع كل سطر بالنقطة) وليست تقديرية.

build_document_lines() تبني قائمة الأسطر النهائية (نفس المصدر تستخدمه
كل من المعاينة داخل الواجهة وملف الـ Word الفعلي، حتى تكون المعاينة
مطابقة تماماً للملف).

Commission 1/2 وFrais وTaxe وMotif وDevise: قيم ثابتة زي الأصل بالضبط
(مو حقول تُملأ)، بانتظار تحديد غير ذلك لاحقاً.
"""
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Mm
from PIL import Image, ImageDraw, ImageFont

# معرَّفة بـui/common/widgets.py (خانات التاريخ هي المستخدم الأساسي
# الحقيقي — 7 استخدامات مقابل استخدام وحيد هنا) — نفس الأسماء المستخدمة
# بالنموذج الأصلي (بدون همزات/أكسنت: "Aout" لا "Août").
from ui.common.widgets import FRENCH_MONTHS
from programme.paths import get_screen_dir

# مكان حفظ مستندات CD الحقيقية: مجلد travail/CD بسطح المكتب — مو جوا
# مجلد البرنامج نفسه (راجع paths.py للتفاصيل والسبب).
OUTPUT_DIR = get_screen_dir("CD")

# --- قياسات مستخرجة رقمياً من Model Change Devise.pdf وModel Vierge.pdf عبر PyMuPDF ---
FONT_NAME = "Courier New"
FONT_SIZE = 10  # العنوان + التاريخ/الوقت + الصندوق تحت (زي Model Vierge بالضبط)
FONT_SIZE_MIDDLE = 11  # من Agence لين Net a créditer (زي Model Vierge بالضبط)
CHAR_WIDTH_PT = 6.0  # عرض الحرف الواحد بخط Courier New 10pt (0.6em)
LEFT_MARGIN_PT = 30.5
TOP_MARGIN_PT = 58  # يوضع العنوان (أول سطر) عند نفس ارتفاع الأصل تقريباً
LINE_HEIGHT_PT = 11.3  # تباعد أسطر خط 10
LINE_HEIGHT_MIDDLE_PT = 12.45  # تباعد أسطر خط 11 (من Model Vierge)

TITLE_COL = 25   # عدد المسافات قبل العنوان (محسوبة من: 182.8-30.5 / 6.0)
DATE_COL = 54    # عدد المسافات قبل التاريخ (محسوبة من: 354.5-30.5 / 6.0)

# الأسطر الخمسة الأولى (Agence/Devise/Guichet/Caisse/Guichetier) كلها
# بنفس النمط الموحّد (بطلب صريح): حقل أول 5 خانات (أرقام إلا Devise:
# حروف كبيرة فقط)، فراغ واحد ثابت، ثم حقل ثاني 25 حرف (حروف كبيرة أو
# أرقام، بلا مسافة) — ما عدا Guichetier اللي ما فيه حقل ثاني إطلاقاً
# (رقم 5 خانات بس، لوحده على السطر). عمود الحقل الثاني بكل سطر ثابت
# دائماً بغض النظر عن طول الحقل الأول الفعلي المكتوب (.ljust() بالعمود
# الكامل، لا بطول القيمة) — نفس مبدأ Guichet قبل اسم الراكب المكرر.
FIRST_FIELD_WIDTH = 5
FIRST_FIELD_GAP = 1
SECOND_FIELD_WIDTH = 25

AGENCE_LABEL = "Agence ......: "
AGENCE_NO_COL = len(AGENCE_LABEL)  # = 15
AGENCE_NO_FIELD_WIDTH = FIRST_FIELD_WIDTH
AGENCE_GAP = FIRST_FIELD_GAP
AGENCE_COL = AGENCE_NO_COL + AGENCE_NO_FIELD_WIDTH + AGENCE_GAP  # = 21
AGENCE_FIELD_WIDTH = SECOND_FIELD_WIDTH

DEVISE_LABEL = "Devise ......: "
DEVISE_CODE_COL = len(DEVISE_LABEL)  # = 15
DEVISE_CODE_FIELD_WIDTH = FIRST_FIELD_WIDTH
DEVISE_GAP = FIRST_FIELD_GAP
DEVISE_COL = DEVISE_CODE_COL + DEVISE_CODE_FIELD_WIDTH + DEVISE_GAP  # = 21
DEVISE_FIELD_WIDTH = SECOND_FIELD_WIDTH

GUICHET_LABEL = "Guichet .....: "
GUICHET_NO_COL = len(GUICHET_LABEL)  # = 15
GUICHET_NO_FIELD_WIDTH = FIRST_FIELD_WIDTH
GUICHET_GAP = FIRST_FIELD_GAP
GUICHET_COL = GUICHET_NO_COL + GUICHET_NO_FIELD_WIDTH + GUICHET_GAP  # = 21
GUICHET_FIELD_WIDTH = SECOND_FIELD_WIDTH

# مكان تكرار اسم الراكب مثبّت هنا نهائياً — بطلب صريح ما ينزاح أبداً حتى
# بعد تقسيم حقل Guichet لحقلين (5+1+25 = 31 خانة، أقل من الـ32 المحجوزة
# هنا، فيبقى فراغ حرف واحد كحاجز أمان قبل بداية التكرار، بلا أي تصادم).
# +32 من GUICHET_NO_COL (لا +33): GUICHET_NO_COL نفسه أصلاً بعد الفاصل
# الأول (فراغ واحد بعد النقطتين)، فـ33 فراغ الإجمالية من النقطتين =
# 1(الفاصل، مُحتسَب أصلاً بـGUICHET_NO_COL) + 32 بعده.
GUICHET_NAME_COL = GUICHET_NO_COL + 32

CAISSE_LABEL = "Caisse ......: "
CAISSE_NO_COL = len(CAISSE_LABEL)  # = 15
CAISSE_NO_FIELD_WIDTH = FIRST_FIELD_WIDTH
CAISSE_GAP = FIRST_FIELD_GAP
CAISSE_COL = CAISSE_NO_COL + CAISSE_NO_FIELD_WIDTH + CAISSE_GAP  # = 21
CAISSE_FIELD_WIDTH = SECOND_FIELD_WIDTH

# Guichetier استثناء: حقل واحد بس (رقم 5 خانات)، بلا حقل ثاني إطلاقاً.
GUICHETIER_LABEL = "Guichetier ..: "
GUICHETIER_COL = len(GUICHETIER_LABEL)  # = 15
GUICHETIER_FIELD_WIDTH = FIRST_FIELD_WIDTH

# سطر Nature piece identite: فراغ واحد بعد النقطتين، "PSP" ثابت، 8
# فراغات ثابتة، "No" ثابت، فراغ، نقطتين — وبعد هالنقطتين مباشرة (بلا
# فراغ) يبدأ حقل N° Passport (12 حرف/رقم بالضبط، هذا حده الأقصى). بعده
# 5 فراغات ثابتة، "Obtent." ثابت، فراغ، نقطتين — وبعد هالنقطتين مباشرة
# (بلا فراغ) يبدأ حقل تاريخ الحصول عليها (DD/MM/YYYY، بفواصل "/" دائماً).
# (قِيست هذي الأعداد بدقة من ملف Model Change Devise.pdf نفسه — استخراج
# موضع كل حرف بالنقطة عبر PyMuPDF — بعد ما تبيّن إن نسخة سابقة كانت
# مبنية على عدّ يدوي غلط: 10 فراغات لا 8، و"No:" بلا فراغ قبل النقطتين
# بدل "No :" الصحيحة، فرق عمود واحد كان يخلي حقل Obtent وTx de change
# ينزاحان عن مكانهما الحقيقي.)
PASSPORT_LABEL = "Nature piece identite : PSP" + " " * 8 + "No :"
PASSPORT_VALUE_COL = len(PASSPORT_LABEL)
PASSPORT_FIELD_WIDTH = 12
OBTENT_GAP = 5  # فراغات ثابتة بين نهاية حقل N° Passport وبداية "Obtent."
OBTENT_LABEL = "Obtent. :"
OBTENT_COL = PASSPORT_VALUE_COL + PASSPORT_FIELD_WIDTH + OBTENT_GAP
DATE_DELIVRANCE_COL = OBTENT_COL + len(OBTENT_LABEL)  # مباشرة بعد النقطتين، بلا فراغ

# Motif ثابت دائماً حسب طلبك (Commission/Frais/Taxe كمان دائماً 0 بالأسفل).
# باقي الخانات الفاضية (No، Devise، N° Passport...) تُعبّى من لوحة التحكم؛
# ما فيها قيمة افتراضية ولا نقاط تعبئة — فاضية لين تُملأ.
MOTIF_TEXT = "Cession DEV nationaux résident"

# سطر Montant en devise: حقل كتابة المبلغ (اليورو) لازم ينتهي دائماً بفاصل
# مسافة واحدة بالضبط قبل كلمة "EUR" — هالمكان ثابت ما يتحرك أبداً (نفس
# مكانه الأصلي). لكن عرض الحقل الفعلي المسموح للكتابة فيه صار 10 أرقام
# بس (بدل 20)، فبداية الحقل تقدّمت 10 خانات نحو اليمين، وباقي كل شي —
# التسمية، النقاط، "EUR"، "Tx de change" — بمكانه الأصلي بدون أي تغيير.
EUR_TOTAL_SLOT = 20   # العرض الإجمالي الأصلي المحجوز قبل " EUR" (ثابت، ما يتحرك)
EUR_FIELD_WIDTH = 10  # عرض حقل الكتابة الفعلي الجديد (10 أرقام بالضبط)
EUR_COL = 24 + (EUR_TOTAL_SLOT - EUR_FIELD_WIDTH)  # = 34

# حقل Tx de change (taux): بعد النقطتين — 6 فراغات ثابتة، ثم حقل الكتابة.
# صيغة فرنسية (فاصلة عشرية "," زي باقي حقول المبالغ) بس 3 أرقام صحيحة
# بالضبط كحد أقصى (999) و7 أرقام عشرية بالضبط دائماً (بدل 2 مثل EUR/DZD)
# — يعني السعة الإجمالية 3 + فاصلة + 7 = 11 خانة، والرقم دائماً أقصى
# اليمين (مسافات فاضية على اليسار لو الجزء الصحيح أقل من 3 أرقام).
TX_DE_CHANGE_LABEL = "Tx de change :"
TAUX_GAP = 6
TAUX_INT_DIGITS = 3
TAUX_DEC_DIGITS = 7
TAUX_FIELD_WIDTH = TAUX_INT_DIGITS + 1 + TAUX_DEC_DIGITS  # = 11 (3 + فاصلة + 7)
TAUX_MAX_VALUE = 10 ** TAUX_INT_DIGITS - 1  # = 999


def _french_date(d):
    return f"{d.day} {FRENCH_MONTHS[d.month - 1]}"


def _date_part_text(d):
    """نص التاريخ زي ما يُكتب بالسطر (يوم + شهر بالفرنسي + سنة)، أو فاضي
    لو التاريخ None — مصدر واحد يستخدمه كل من سطر التاريخ الحقيقي وحساب
    مكان حقل الوقت، حتى ما ينفصلوا أبداً."""
    return f"{_french_date(d)}  {d.year}" if d else ""


def _pad(n):
    return " " * n


def _fr_amount(value, decimals=2):
    """تنسيق فرنسي/جزائري: فاصلة عشرية "," وفاصل آلاف "." (زي 24.177,60)."""
    s = f"{value:,.{decimals}f}"
    return s.translate(str.maketrans(",.", ".,"))


def _value_field(value_str, width=20):
    return f"{value_str.rjust(width)} DZD"


def _label(text, width=22):
    """يبني تسمية بنقط تعبئة تلقائية بعرض ثابت، حتى تصطف كل النقطتين
    (":") ببعضها بالضبط زي الأصل — بدل عدّ النقاط يدوياً وغلط فيها."""
    return f"{text} ".ljust(width, ".") + ":"


# عمود بداية حقل "Nom du remettant" محسوب مباشرة من نفس التسمية اللي
# تُبنى بيها السطر الحقيقي (_label) + فراغ واحد بعد النقطتين — مو رقم
# ثابت مكتوب يدوياً، حتى يستحيل ينحرف حقل الكتابة الحي عن مكانه الفعلي
# بالمستند حتى لو تغيّر نص التسمية أو عرضها لاحقاً.
NOM_REMETTANT_COL = len(_label("Nom du remettant")) + 1

# عمود بداية حقل taux محسوب من نفس الأجزاء الثابتة اللي تُبنى بيها السطر
# الحقيقي (تسمية Montant en devise + حقل EUR بعرضه الإجمالي + " EUR " +
# تسمية Tx de change) — مو رقم ثابت مكتوب يدوياً.
_MONTANT_PREFIX_LEN = len(_label("Montant en devise")) + 1 + EUR_TOTAL_SLOT + len(f" EUR {TX_DE_CHANGE_LABEL}")
TAUX_COL = _MONTANT_PREFIX_LEN + TAUX_GAP


def _mono_paragraph(doc, text="", size=FONT_SIZE):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    line_height = LINE_HEIGHT_MIDDLE_PT if size == FONT_SIZE_MIDDLE else LINE_HEIGHT_PT
    p.paragraph_format.line_spacing = Pt(line_height)  # تباعد "ثابت" مطابق للأصل
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = False
    return p


def _build_document_entries(data):
    """
    data: dict فيه no, date, time, agence, guichet, caisse, guichetier,
    passager, passport_no, date_delivrance, taux, eur, dzd

    يرجّع قائمة (نص، حجم الخط) جاهزة (فيها كل المسافات محسوبة مسبقاً)، كل
    سطر نص واحد يُكتب كما هو بمحاذاة يسار (زي الأصل بالضبط).
    """
    no_text = data["no"]
    delivrance_str = data["date_delivrance"].strftime("%d/%m/%Y") if data["date_delivrance"] else ""

    # Taux وEUR: فاضيين لو لوحة التحكم ما عبّتهم (بدل ما نحط 0 افتراضياً).
    # taux بصيغة فرنسية (فاصلة عشرية "،" زي باقي حقول المبالغ) بـ7 أرقام
    # عشرية بالضبط دائماً، محاذى أقصى اليمين بعرض حقله الكامل (11 خانة).
    taux_part = (
        _fr_amount(data["taux"], decimals=TAUX_DEC_DIGITS).rjust(TAUX_FIELD_WIDTH)
        if data["taux"] is not None else " " * TAUX_FIELD_WIDTH
    )
    # rjust بعرض الـslot الإجمالي (20) — مو عرض حقل الكتابة (10) — حتى
    # "EUR" يبقى بنفس مكانه الثابت دائماً بغض النظر عن عرض حقل الكتابة.
    eur_part = _fr_amount(data['eur']).rjust(EUR_TOTAL_SLOT) if data["eur"] is not None else " " * EUR_TOTAL_SLOT
    # Soit وNet a créditer: فاضيين بردو لو DZD ما تعبّى (بدل 0,00 ثابتة)،
    # نفس مبدأ باقي الحقول اللي تحددها لوحة التحكم.
    dzd_value_str = _fr_amount(data["dzd"]) if data["dzd"] is not None else ""

    # حرف "a" مو بعمود ثابت — مكانه يتحدد بطول التاريخ لي قبلو (يوم + شهر
    # + سنة)، بس دائماً بفاصل مسافة واحدة بالضبط بعده، بغض النظر عن طول
    # اسم الشهر (حتى لو "Septembre" الطويلة) أو حتى لو التاريخ فاضي كلياً.
    # (خلفية الشاشة الحية فقط: data["_omit_a_literal"] يشيل "a" من الصورة
    # المرسومة، لأن حرف a هناك يُرسم حياً فوقها بشاشة العمل، مو مطبوع
    # ثابت بالخلفية — حتى ما يصير ظهور مزدوج/مضبب. المستند الحقيقي دايماً
    # يحتوي "a" عادي لأن هالمفتاح ما يوصلها أبداً.)
    time_part = data["time"] or ""
    date_part = _date_part_text(data["date"])
    a_char = "" if data.get("_omit_a_literal") else "a"
    date_line = _pad(DATE_COL) + date_part + f" {a_char} {time_part}"

    S = FONT_SIZE          # قسم العنوان/التاريخ/الصندوق: خط 10
    M = FONT_SIZE_MIDDLE   # قسم Agence لين Net a créditer: خط 11

    entries = [
        (f"{_pad(TITLE_COL)}BORDEREAU D'ACHAT DEVISE No  {no_text}", S),
        ("", S),
        ("", S),
        (date_line, S),
        ("", S),
        ("", S),
        ("", S),
        ("", S),
        (f"{AGENCE_LABEL}{data.get('agence_no', '')}".ljust(AGENCE_COL) + data['agence'], M),
        (f"{DEVISE_LABEL}{data.get('devise_code', '')}".ljust(DEVISE_COL) + data.get('devise', ''), M),
        (
            (f"{GUICHET_LABEL}{data.get('guichet_no', '')}".ljust(GUICHET_COL) + data['guichet']).ljust(GUICHET_NAME_COL)
            + data['passager'],
            M,
        ),
        (f"{CAISSE_LABEL}{data.get('caisse_no', '')}".ljust(CAISSE_COL) + data['caisse'], M),
        (f"{GUICHETIER_LABEL}{data['guichetier']}", M),
        ("", M),
        ("", M),
        (f"{_label('Nom du remettant')} {data['passager']}", M),
        (
            f"{PASSPORT_LABEL}{data['passport_no'].ljust(PASSPORT_FIELD_WIDTH)}"
            + " " * OBTENT_GAP
            + f"{OBTENT_LABEL}{delivrance_str}",
            M,
        ),
        (f"{_label('Motif')} {MOTIF_TEXT}", M),
        (
            f"{_label('Montant en devise')} {eur_part} EUR {TX_DE_CHANGE_LABEL}"
            + " " * TAUX_GAP
            + taux_part,
            M,
        ),
        (f"{_label('Soit')} {_value_field(dzd_value_str)}", M),
        (f"{_label('Commission  1')} {_value_field(_fr_amount(0))}", M),
        (f"{_label('Commission  2')} {_value_field(_fr_amount(0))}", M),
        (f"{_label('Frais')} {_value_field(_fr_amount(0))}", M),
        (f"{_label('Taxe')} {_value_field(_fr_amount(0))}", M),
        (f"{_label('Net a créditer')} {_value_field(dzd_value_str)}", M),
    ]
    entries += [("", S)] * 9
    entries += [
        (" --------------------------", S),
        ("!   CLIENT   !  GUICHETIER !", S),
        ("!            !             !", S),
        ("!            !             !", S),
        ("!            !             !", S),
        ("!            !             !                            OPERATION EFFECTUEE", S),
    ]
    return entries


def build_document_lines(data):
    """نص فقط بدون حجم الخط (تُستخدم بالاختبارات/المقارنة النصية)."""
    return [text for text, _size in _build_document_entries(data)]


def _build_docx(data):
    """يبني مستند Word كامل (بدون حفظ) من بيانات النموذج."""
    entries = _build_document_entries(data)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)   # عرض A4 (بدون تغيير)
    section.page_height = Mm(203)  # طول مخصّص أقصر من A4 حسب طلبك (20.3 سم)
    section.left_margin = Pt(LEFT_MARGIN_PT)
    section.right_margin = Pt(LEFT_MARGIN_PT)
    section.top_margin = Pt(TOP_MARGIN_PT)
    section.bottom_margin = Pt(LEFT_MARGIN_PT)

    for text, size in entries:
        _mono_paragraph(doc, text, size)

    return doc


def _named_path(dir_path, passager, ext):
    """مسار ملف جديد داخل مجلد محدَّد صراحة، بنفس اصطلاح التسمية دائماً
    (CD_<الراكب>_<الوقت>.<الامتداد>) — يُستخدم لكل من المجلد التلقائي
    (مجلد الشهر، راجع _month_output_path) ولمكان يختاره المستخدم صراحة
    ("حفظ في مكان آخر" بـtab.py)."""
    os.makedirs(dir_path, exist_ok=True)
    safe_passager = "".join(c for c in passager if c.isalnum() or c in " _-").strip() or "sans_nom"
    filename = f"CD_{safe_passager}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"
    return os.path.join(dir_path, filename)


def _month_output_path(passager, ext):
    """مسار الحفظ التلقائي الافتراضي لأي مستند CD جديد: مجلد فرعي حسب
    الشهر/السنة الحالي (travail/CD/2026-08/...) — أسهل للأرشفة اليدوية
    لاحقاً بدل تراكم كل الملفات بمجلد واحد مسطّح."""
    month_dir = os.path.join(OUTPUT_DIR, datetime.now().strftime("%Y-%m"))
    return _named_path(month_dir, passager, ext)


def generate_cd_document(data, out_path=None):
    """يبني المستند ويحفظه كملف Word نهائي، ويرجّع مساره. out_path
    (اختياري): مسار محدَّد صراحة (تحديث فوق حالة موجودة، أو حفظ بمكان
    آخر يختاره المستخدم) — بلا تمرير، يُبنى تلقائياً بمجلد الشهر الحالي
    (نفس الاصطلاح دائماً)."""
    doc = _build_docx(data)
    if out_path is None:
        out_path = _month_output_path(data["passager"], "docx")
    else:
        os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path


# متوسّط قيمتَي ascent/descent الرسميتَين لخط Courier القياسي بصيغة PDF
# (629 و157 من 1000، عن pdfmetrics.getFont("Courier").face) — نستخدمها
# لحساب موضع خط الأساس (baseline) بمنتصف صندوق ارتفاع كل سطر بالضبط،
# نفس مبدأ anchor="lm" المستخدم لرسم الخلفية بـPillow (راجع
# render_blank_background)، حتى النص يصطف بمنتصف سطره رأسياً لا أعلاه.
_COURIER_PDF_ASCENT = 0.629
_COURIER_PDF_DESCENT = 0.157


def generate_cd_pdf(data, out_path=None):
    """يبني نسخة PDF من نفس بيانات الاستمارة — نص متجهي حقيقي (حاد
    بالطباعة/التكبير، مو صورة)، بنفس مصدر النص الوحيد وحسابات الموضع
    المستخدمة بكل مكان آخر بالمشروع (_build_document_entries، _y_top_pt،
    _line_height_for) — بلا أي اعتماد على Word أو برنامج خارجي (خط
    Courier القياسي مدمج بصيغة PDF نفسها، عرض حرفه 0.6em بالضبط زي
    المستخدم بكل حساباتنا، فما يحتاج أي معايرة إضافية أفقياً)."""
    from reportlab.pdfgen.canvas import Canvas

    if out_path is None:
        out_path = _month_output_path(data["passager"], "pdf")
    else:
        os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)

    c = Canvas(out_path, pagesize=(PAGE_WIDTH_PT, PAGE_HEIGHT_PT))
    for index, (text, size) in enumerate(_build_document_entries(data)):
        if not text.strip():
            continue
        c.setFont("Courier", size)
        line_h = _line_height_for(index)
        glyph_h = (_COURIER_PDF_ASCENT + _COURIER_PDF_DESCENT) * size
        top_pad = (line_h - glyph_h) / 2
        baseline_from_top = _y_top_pt(index) + top_pad + _COURIER_PDF_ASCENT * size
        c.drawString(LEFT_MARGIN_PT, PAGE_HEIGHT_PT - baseline_from_top, text)
    c.showPage()
    c.save()
    return out_path


# --- سجل مواضع الحقول القابلة للتعديل، فوق صورة النموذج مباشرة ---
# كل حقل: (رقم السطر بقائمة _build_document_entries، عمود بداية القيمة
# بالحرف، عرض الخانة بالحروف). حجم الخط يُستنتج من رقم السطر (نفس منطق
# _build_document_entries: 8-24 خط 11، الباقي خط 10).
FIELD_LAYOUT = {
    "no": (0, 54, 12),   # فراغين قبل الرقم أصلاً جزء من النص الأصلي "No  " — حده الأقصى الحقيقي 12 رقم
    "date": (3, 54, 13),
    # عمود "time" هنا (70) قيمة احتياطية غير مستخدمة فعلياً — cd_tab.py
    # يحسب مكانه الحقيقي حياً بالبكسل تبعاً لحرف "a" (راجع SplitDateEntry
    # .a_right_edge_px)، مُستخدم منها بس رقم السطر (3) لحساب الارتفاع y.
    "time": (3, 70, 6),
    "agence_no": (8, AGENCE_NO_COL, AGENCE_NO_FIELD_WIDTH),
    "agence": (8, AGENCE_COL, AGENCE_FIELD_WIDTH),
    "devise_code": (9, DEVISE_CODE_COL, DEVISE_CODE_FIELD_WIDTH),
    "devise": (9, DEVISE_COL, DEVISE_FIELD_WIDTH),
    "guichet_no": (10, GUICHET_NO_COL, GUICHET_NO_FIELD_WIDTH),
    "guichet": (10, GUICHET_COL, GUICHET_FIELD_WIDTH),
    "caisse_no": (11, CAISSE_NO_COL, CAISSE_NO_FIELD_WIDTH),
    "caisse": (11, CAISSE_COL, CAISSE_FIELD_WIDTH),
    "guichetier": (12, GUICHETIER_COL, GUICHETIER_FIELD_WIDTH),
    "passager": (15, NOM_REMETTANT_COL, 32),
    "passport_no": (16, PASSPORT_VALUE_COL, PASSPORT_FIELD_WIDTH),
    "date_delivrance": (16, DATE_DELIVRANCE_COL, 12),
    "eur": (18, EUR_COL, EUR_FIELD_WIDTH),
    "taux": (18, TAUX_COL, TAUX_FIELD_WIDTH),
    # Soit (dzd): نفس عمود وعرض حقل Montant en devise (EUR) بالضبط —
    # الاثنان مشتقّان من نفس الثوابت عمداً، بطلب صريح "نفس السعة
    # والخصائص"، حتى ما ينحرفوا عن بعض لو تغيّرت لاحقاً.
    "dzd": (19, EUR_COL, EUR_FIELD_WIDTH),
    "net_crediter": (24, 24, 20),
    "guichet_mirror": (10, GUICHET_NAME_COL, 32),
}

PAGE_WIDTH_PT = 210 * 2.8346456692913385
PAGE_HEIGHT_PT = 203 * 2.8346456692913385


def _line_font_size(index):
    return FONT_SIZE_MIDDLE if 8 <= index <= 24 else FONT_SIZE


def _line_height_for(index):
    return LINE_HEIGHT_MIDDLE_PT if _line_font_size(index) == FONT_SIZE_MIDDLE else LINE_HEIGHT_PT


def _y_top_pt(index):
    """أعلى نقطة (بالنقطة PDF) لسطر رقمه index، من أعلى الصفحة."""
    return TOP_MARGIN_PT + sum(_line_height_for(i) for i in range(index))


def field_layout_px(target_width_px):
    """
    يرجّع dict: اسم الحقل -> (x, y, width, height) بالبكسل، عند عرض صورة
    مقاسه target_width_px (نفس مقاس الصورة المعروضة بالواجهة).

    ملاحظة: عمود "time" هنا احتياطي فقط — cd_tab.py يحسب مكانه الحقيقي
    حياً بالبكسل من date_entry.a_right_edge_px() (راجع FIELD_LAYOUT).
    """
    scale = target_width_px / PAGE_WIDTH_PT
    result = {}
    for name, (index, col, width_chars) in FIELD_LAYOUT.items():
        size = _line_font_size(index)
        char_w = 0.6 * size
        x_pt = LEFT_MARGIN_PT + col * char_w
        y_pt = _y_top_pt(index)
        w_pt = width_chars * char_w
        h_pt = _line_height_for(index)
        result[name] = (x_pt * scale, y_pt * scale, w_pt * scale, h_pt * scale)
    return result


# --- ملف المعاينة المؤقت (يُستبدل في كل ضغطة "معاينة"، ما يتراكم) ---
# (كان فيه أيضاً PREVIEW_DOCX/PREVIEW_PDF وسيطين لما كان الرسم يمر عبر
# Word — انشالوا مع render_preview_image نفسها، راجع render_blank_background
# فوق: رسم مباشر ببايثون/Pillow، بلا حاجة لـWord ولا لملف PDF وسيط إطلاقاً.)
PREVIEW_PNG = os.path.join(OUTPUT_DIR, "_preview.png")


_EMPTY_DATA = {
    "no": "", "date": None, "time": "", "agence_no": "", "agence": "",
    "devise_code": "", "guichet_no": "", "guichet": "", "caisse_no": "", "caisse": "",
    "guichetier": "", "passager": "", "passport_no": "", "date_delivrance": None,
    "taux": None, "eur": None, "dzd": None,
    "_omit_a_literal": True,  # الخلفية الفاضية بس — راجع الشرح أعلاه
}

# مسار خط Courier New الحقيقي على وندوز — نفس الخط المستخدم أصلاً بكل
# من ملف الـWord (FONT_NAME) وحقول الكتابة الحية بالواجهة، فرسم الخلفية
# به مباشرة (بدل الاعتماد على Word لرسمها) يحافظ على نفس المظهر تماماً.
COURIER_TTF_PATH = r"C:\Windows\Fonts\cour.ttf"

_font_cache = {}


def _courier_font(px_size):
    """يرجّع خط Courier New Pillow بالحجم المطلوب (بالبكسل، كسور مسموحة
    لدقة أعلى) — مع تخزين مؤقت بالذاكرة (تحميل ملف الخط من القرص مكلف
    نسبياً، ونفس الحجم يتكرر لعشرات الأسطر بكل رسمة)."""
    key = round(px_size, 2)
    font = _font_cache.get(key)
    if font is None:
        font = ImageFont.truetype(COURIER_TTF_PATH, px_size)
        _font_cache[key] = font
    return font


def render_blank_background(data, target_width_px=750, out_png=None):
    """
    يرسم صورة المستند مباشرة ببايثون (Pillow) — بدون فتح Word إطلاقاً.
    يستخدم بالضبط نفس مصدر النص الوحيد (_build_document_entries) ونفس
    حسابات الموضع المستخدمة أصلاً لحقول الكتابة الحية (_y_top_pt،
    _line_font_size، وصيغة عرض الحرف 0.6×حجم الخط بـfield_layout_px)،
    فيضمن تطابق الخلفية مع الحقول الحية تماماً بلا أي اعتماد على تثبيت
    Word أو تشغيله — أسرع بكثير كمان (رسم مباشر، بلا فتح برنامج ولا
    تحويل PDF وسيط).
    """
    out_png = out_png or PREVIEW_PNG
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    scale = target_width_px / PAGE_WIDTH_PT
    img_h = round(PAGE_HEIGHT_PT * scale)
    img = Image.new("RGB", (target_width_px, img_h), "white")
    draw = ImageDraw.Draw(img)

    for index, (text, size) in enumerate(_build_document_entries(data)):
        if not text.strip():
            continue  # سطر فاضي كلياً — ما فيه شي نرسمه
        font = _courier_font(size * scale)
        char_w_px = 0.6 * size * scale
        # نرسم النص بمنتصف صندوق ارتفاع السطر رأسياً (anchor="lm") — نفس
        # المبدأ اللي يعتمده Tk تلقائياً بخانة الكتابة الحية المفروض
        # عليها نفس الارتفاع (h من field_layout_px)، فالنص الثابت يصطف
        # على نفس السطر بالضبط مع أي حقل حي جنبه بغض النظر عن أبعاد
        # الخط الداخلية (ascent/descent) — بدل تخمين هامش علوي يدوياً.
        y_center_px = (_y_top_pt(index) + _line_height_for(index) / 2) * scale
        # نرسم كل حرف لحاله بعموده المحسوب مباشرة (مو السطر كامل بدفعة
        # وحدة) — تأكدنا (بالقياس المباشر) إن FreeType يطبّق تصحيح تقريب
        # (hinting) على تقدّم كل حرف عند أحجام الخط الصغيرة (زي حجمنا
        # هنا)، فيتراكم انحراف ملحوظ عبر سطر طويل (وصل ~24px بسطر 54
        # حرف) لو رسمناه كنص واحد متصل. الرسم حرف-حرف بموضعه الدقيق
        # يتجاوز هالمشكلة كلياً ويضمن مطابقة تامة مع أعمدة field_layout_px.
        for col, ch in enumerate(text):
            if ch == " ":
                continue
            x_px = LEFT_MARGIN_PT * scale + col * char_w_px
            draw.text((x_px, y_center_px), ch, font=font, fill="black", anchor="lm")

    img.save(out_png)
    return out_png


def get_blank_background(target_width_px=750, force=False):
    """
    يرسم صورة النموذج الفاضي (بس الأشياء الثابتة: Motif، PSP، Commission/
    Frais/Taxe = 0، الصندوق تحت) بنفس الدقة المطلوبة (target_width_px)
    ويخزّنها بملف خاص بهالدقة — حتى تكون الكتابة حادة دائماً بأي مستوى
    زوم (رسم أصلي بكل دقة، مو تكبير صورة صغيرة فتصير مبكسلة/مضببة).
    كل دقة تتولّد مرة وحدة بس (أول استخدام لها)، وتُعاد استخدامها بعدها.
    """
    path = os.path.join(OUTPUT_DIR, f"_blank_bg_{target_width_px}.png")
    if not force and os.path.exists(path):
        return path
    return render_blank_background(_EMPTY_DATA, target_width_px=target_width_px, out_png=path)
