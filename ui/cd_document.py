"""
بناء وتوليد مستند "Change Devise" بصيغة Word، مطابق حرفياً (خط، حجم،
هوامش، مسافات، تنسيق الأرقام) لنموذج البوردرو الأصلي (Model Change
Devise.pdf) — القياسات مستخرجة رقمياً من ملف الـ PDF نفسه (نوع الخط،
حجمه، وموضع كل سطر بالنقطة) وليست تقديرية.

build_document_lines() تبني قائمة الأسطر النهائية (نفس المصدر تستخدمه
كل من المعاينة داخل الواجهة وملف الـ Word الفعلي، حتى تكون المعاينة
مطابقة تماماً للملف).

Commission 1/2 وFrais وTaxe وMotif وDevise: قيم ثابتة زي الأصل بالضبط
(مو حقول تُملأ)، بانتظار تحديد غير ذلك لاحقاً.
"""
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Mm

# نفس الأسماء المستخدمة بالنموذج الأصلي (بدون همزات/أكسنت، زي ما هي
# بالمستند المصدر: "Aout" لا "Août").
FRENCH_MONTHS = [
    "Janvier", "Fevrier", "Mars", "Avril", "Mai", "Juin",
    "Juillet", "Aout", "Septembre", "Octobre", "Novembre", "Decembre",
]

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "output", "CD")

# --- قياسات مستخرجة رقمياً من Model Change Devise.pdf وModel Vierge.pdf عبر PyMuPDF ---
FONT_NAME = "Courier New"
FONT_SIZE = 10  # العنوان + التاريخ/الوقت + الصندوق تحت (زي Model Vierge بالضبط)
FONT_SIZE_MIDDLE = 11  # من Agence لين Net a créditer (زي Model Vierge بالضبط)
CHAR_WIDTH_PT = 6.0  # عرض الحرف الواحد بخط Courier New 10pt (0.6em)
LEFT_MARGIN_PT = 30.5
TOP_MARGIN_PT = 58  # يوضع العنوان (أول سطر) عند نفس ارتفاع الأصل تقريباً
LINE_HEIGHT_PT = 11.3  # تباعد أسطر خط 10
LINE_HEIGHT_MIDDLE_PT = 12.45  # تباعد أسطر خط 11 (من Model Vierge)

TITLE_COL = 25   # عدد المسافات قبل العنوان (محسوبة من: 182.8-30.5 / 6.0)
DATE_COL = 54    # عدد المسافات قبل التاريخ (محسوبة من: 354.5-30.5 / 6.0)

# سطر Guichet: فراغ واحد بعد النقطتين، ثم حقل الكتابة (30 حرف بالضبط —
# هذا حده الأقصى)، ثم فراغين، ثم يبدأ اسم الراكب المكرر — يعني 33 فراغ
# بالضبط من النقطتين لبداية التكرار (1 فاصل + 30 حقل + 2 فراغ).
GUICHET_LABEL = "Guichet .....: "
GUICHET_VALUE_COL = len(GUICHET_LABEL)  # فراغ واحد بعد النقطتين
GUICHET_FIELD_WIDTH = 30                # سعة حقل الكتابة بالضبط الحالية

# مكان التكرار مثبّت هنا نهائياً — بطلب صريح ما ينزاح أبداً حتى لو تغيّر
# عرض حقل الكتابة (GUICHET_FIELD_WIDTH) لاحقاً. الفراغ الفاصل بينهما هو
# اللي يتمدد أو يضيق تلقائياً بدل ما يتحرك مكان التكرار نفسه.
# +32 من GUICHET_VALUE_COL (لا +33): GUICHET_VALUE_COL نفسه أصلاً بعد
# الفاصل الأول (فراغ واحد بعد النقطتين)، فـ33 فراغ الإجمالية من النقطتين
# = 1(الفاصل، مُحتسَب أصلاً بـGUICHET_VALUE_COL) + 32 بعده.
GUICHET_NAME_COL = GUICHET_VALUE_COL + 32

# سطر Nature piece identite: فراغ واحد بعد النقطتين، "PSP" ثابت، 10
# فراغات ثابتة، "No" ثابت، نقطتين مباشرة (بلا فراغ بينهما، تحقّقنا من
# الأصل) — وبعد هالنقطتين مباشرة كمان (بلا فراغ) يبدأ حقل N° Passport
# (12 حرف/رقم بالضبط، هذا حده الأقصى). بعده 5 فراغات ثابتة، "Obtent."
# ثابت، فراغ، نقطتين — وبعد هالنقطتين مباشرة (بلا فراغ) يبدأ حقل تاريخ
# الحصول عليها (DD/MM/YYYY، بفواصل "/" دائماً).
PASSPORT_LABEL = "Nature piece identite : PSP" + " " * 10 + "No:"
PASSPORT_VALUE_COL = len(PASSPORT_LABEL)
PASSPORT_FIELD_WIDTH = 12
OBTENT_GAP = 5  # فراغات ثابتة بين نهاية حقل N° Passport وبداية "Obtent."
OBTENT_LABEL = "Obtent. :"
OBTENT_COL = PASSPORT_VALUE_COL + PASSPORT_FIELD_WIDTH + OBTENT_GAP
DATE_DELIVRANCE_COL = OBTENT_COL + len(OBTENT_LABEL)  # مباشرة بعد النقطتين، بلا فراغ

# Motif ثابت دائماً حسب طلبك (Commission/Frais/Taxe كمان دائماً 0 بالأسفل).
# باقي الخانات الفاضية (No، Devise، N° Passport...) تُعبّى من لوحة التحكم؛
# ما فيها قيمة افتراضية ولا نقاط تعبئة — فاضية لين تُملأ.
MOTIF_TEXT = "Cession DEV nationaux résident"

# سطر Montant en devise: حقل كتابة المبلغ (اليورو) لازم ينتهي دائماً بفاصل
# مسافة واحدة بالضبط قبل كلمة "EUR" — هالمكان ثابت ما يتحرك أبداً (نفس
# مكانه الأصلي). لكن عرض الحقل الفعلي المسموح للكتابة فيه صار 10 أرقام
# بس (بدل 20)، فبداية الحقل تقدّمت 10 خانات نحو اليمين، وباقي كل شي —
# التسمية، النقاط، "EUR"، "Tx de change" — بمكانه الأصلي بدون أي تغيير.
EUR_TOTAL_SLOT = 20   # العرض الإجمالي الأصلي المحجوز قبل " EUR" (ثابت، ما يتحرك)
EUR_FIELD_WIDTH = 10  # عرض حقل الكتابة الفعلي الجديد (10 أرقام بالضبط)
EUR_COL = 24 + (EUR_TOTAL_SLOT - EUR_FIELD_WIDTH)  # = 34

# حقل Tx de change (taux): بعد النقطتين — 6 فراغات ثابتة، ثم حقل الكتابة.
# صيغة فرنسية (فاصلة عشرية "," زي باقي حقول المبالغ) بس 3 أرقام صحيحة
# بالضبط كحد أقصى (999) و7 أرقام عشرية بالضبط دائماً (بدل 2 مثل EUR/DZD)
# — يعني السعة الإجمالية 3 + فاصلة + 7 = 11 خانة، والرقم دائماً أقصى
# اليمين (مسافات فاضية على اليسار لو الجزء الصحيح أقل من 3 أرقام).
TX_DE_CHANGE_LABEL = "Tx de change :"
TAUX_GAP = 6
TAUX_INT_DIGITS = 3
TAUX_DEC_DIGITS = 7
TAUX_FIELD_WIDTH = TAUX_INT_DIGITS + 1 + TAUX_DEC_DIGITS  # = 11 (3 + فاصلة + 7)
TAUX_MAX_VALUE = 10 ** TAUX_INT_DIGITS - 1  # = 999


def _french_date(d):
    return f"{d.day} {FRENCH_MONTHS[d.month - 1]}"


def _date_part_text(d):
    """نص التاريخ زي ما يُكتب بالسطر (يوم + شهر بالفرنسي + سنة)، أو فاضي
    لو التاريخ None — مصدر واحد يستخدمه كل من سطر التاريخ الحقيقي وحساب
    مكان حقل الوقت، حتى ما ينفصلوا أبداً."""
    return f"{_french_date(d)}  {d.year}" if d else ""


def _pad(n):
    return " " * n


def _fr_amount(value, decimals=2):
    """تنسيق فرنسي/جزائري: فاصلة عشرية "," وفاصل آلاف "." (زي 24.177,60)."""
    s = f"{value:,.{decimals}f}"
    return s.translate(str.maketrans(",.", ".,"))


def _value_field(value_str, width=20):
    return f"{value_str.rjust(width)} DZD"


def _label(text, width=22):
    """يبني تسمية بنقط تعبئة تلقائية بعرض ثابت، حتى تصطف كل النقطتين
    (":") ببعضها بالضبط زي الأصل — بدل عدّ النقاط يدوياً وغلط فيها."""
    return f"{text} ".ljust(width, ".") + ":"


# عمود بداية حقل "Nom du remettant" محسوب مباشرة من نفس التسمية اللي
# تُبنى بيها السطر الحقيقي (_label) + فراغ واحد بعد النقطتين — مو رقم
# ثابت مكتوب يدوياً، حتى يستحيل ينحرف حقل الكتابة الحي عن مكانه الفعلي
# بالمستند حتى لو تغيّر نص التسمية أو عرضها لاحقاً.
NOM_REMETTANT_COL = len(_label("Nom du remettant")) + 1

# عمود بداية حقل taux محسوب من نفس الأجزاء الثابتة اللي تُبنى بيها السطر
# الحقيقي (تسمية Montant en devise + حقل EUR بعرضه الإجمالي + " EUR " +
# تسمية Tx de change) — مو رقم ثابت مكتوب يدوياً.
_MONTANT_PREFIX_LEN = len(_label("Montant en devise")) + 1 + EUR_TOTAL_SLOT + len(f" EUR {TX_DE_CHANGE_LABEL}")
TAUX_COL = _MONTANT_PREFIX_LEN + TAUX_GAP


def _mono_paragraph(doc, text="", size=FONT_SIZE):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    line_height = LINE_HEIGHT_MIDDLE_PT if size == FONT_SIZE_MIDDLE else LINE_HEIGHT_PT
    p.paragraph_format.line_spacing = Pt(line_height)  # تباعد "ثابت" مطابق للأصل
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = False
    return p


def _build_document_entries(data):
    """
    data: dict فيه no, date, time, agence, guichet, caisse, guichetier,
    passager, passport_no, date_delivrance, taux, eur, dzd

    يرجّع قائمة (نص، حجم الخط) جاهزة (فيها كل المسافات محسوبة مسبقاً)، كل
    سطر نص واحد يُكتب كما هو بمحاذاة يسار (زي الأصل بالضبط).
    """
    no_text = data["no"]
    delivrance_str = data["date_delivrance"].strftime("%d/%m/%Y") if data["date_delivrance"] else ""

    # Taux وEUR: فاضيين لو لوحة التحكم ما عبّتهم (بدل ما نحط 0 افتراضياً).
    # taux بصيغة فرنسية (فاصلة عشرية "،" زي باقي حقول المبالغ) بـ7 أرقام
    # عشرية بالضبط دائماً، محاذى أقصى اليمين بعرض حقله الكامل (11 خانة).
    taux_part = (
        _fr_amount(data["taux"], decimals=TAUX_DEC_DIGITS).rjust(TAUX_FIELD_WIDTH)
        if data["taux"] is not None else " " * TAUX_FIELD_WIDTH
    )
    # rjust بعرض الـslot الإجمالي (20) — مو عرض حقل الكتابة (10) — حتى
    # "EUR" يبقى بنفس مكانه الثابت دائماً بغض النظر عن عرض حقل الكتابة.
    eur_part = _fr_amount(data['eur']).rjust(EUR_TOTAL_SLOT) if data["eur"] is not None else " " * EUR_TOTAL_SLOT
    # Soit وNet a créditer: فاضيين بردو لو DZD ما تعبّى (بدل 0,00 ثابتة)،
    # نفس مبدأ باقي الحقول اللي تحددها لوحة التحكم.
    dzd_value_str = _fr_amount(data["dzd"]) if data["dzd"] is not None else ""

    # حرف "a" مو بعمود ثابت — مكانه يتحدد بطول التاريخ لي قبلو (يوم + شهر
    # + سنة)، بس دائماً بفاصل مسافة واحدة بالضبط بعده، بغض النظر عن طول
    # اسم الشهر (حتى لو "Septembre" الطويلة) أو حتى لو التاريخ فاضي كلياً.
    # (خلفية الشاشة الحية فقط: data["_omit_a_literal"] يشيل "a" من الصورة
    # المرسومة، لأن حرف a هناك يُرسم حياً فوقها بشاشة العمل، مو مطبوع
    # ثابت بالخلفية — حتى ما يصير ظهور مزدوج/مضبب. المستند الحقيقي دايماً
    # يحتوي "a" عادي لأن هالمفتاح ما يوصلها أبداً.)
    time_part = data["time"] or ""
    date_part = _date_part_text(data["date"])
    a_char = "" if data.get("_omit_a_literal") else "a"
    date_line = _pad(DATE_COL) + date_part + f" {a_char} {time_part}"

    S = FONT_SIZE          # قسم العنوان/التاريخ/الصندوق: خط 10
    M = FONT_SIZE_MIDDLE   # قسم Agence لين Net a créditer: خط 11

    entries = [
        (f"{_pad(TITLE_COL)}BORDEREAU D'ACHAT DEVISE No  {no_text}", S),
        ("", S),
        ("", S),
        (date_line, S),
        ("", S),
        ("", S),
        ("", S),
        ("", S),
        (f"Agence ......: {data['agence']}", M),
        (f"Devise ......: {data.get('devise', '')}", M),
        (f"{GUICHET_LABEL}{data['guichet']}".ljust(GUICHET_NAME_COL) + data['passager'], M),
        (f"Caisse ......: {data['caisse']}", M),
        (f"Guichetier ..: {data['guichetier']}", M),
        ("", M),
        ("", M),
        (f"{_label('Nom du remettant')} {data['passager']}", M),
        (
            f"{PASSPORT_LABEL}{data['passport_no'].ljust(PASSPORT_FIELD_WIDTH)}"
            + " " * OBTENT_GAP
            + f"{OBTENT_LABEL}{delivrance_str}",
            M,
        ),
        (f"{_label('Motif')} {MOTIF_TEXT}", M),
        (
            f"{_label('Montant en devise')} {eur_part} EUR {TX_DE_CHANGE_LABEL}"
            + " " * TAUX_GAP
            + taux_part,
            M,
        ),
        (f"{_label('Soit')} {_value_field(dzd_value_str)}", M),
        (f"{_label('Commission  1')} {_value_field(_fr_amount(0))}", M),
        (f"{_label('Commission  2')} {_value_field(_fr_amount(0))}", M),
        (f"{_label('Frais')} {_value_field(_fr_amount(0))}", M),
        (f"{_label('Taxe')} {_value_field(_fr_amount(0))}", M),
        (f"{_label('Net a créditer')} {_value_field(dzd_value_str)}", M),
    ]
    entries += [("", S)] * 9
    entries += [
        (" --------------------------", S),
        ("!   CLIENT   !  GUICHETIER !", S),
        ("!            !             !", S),
        ("!            !             !", S),
        ("!            !             !", S),
        ("!            !             !                            OPERATION EFFECTUEE", S),
    ]
    return entries


def build_document_lines(data):
    """نص فقط بدون حجم الخط (تُستخدم بالاختبارات/المقارنة النصية)."""
    return [text for text, _size in _build_document_entries(data)]


def _build_docx(data):
    """يبني مستند Word كامل (بدون حفظ) من بيانات النموذج."""
    entries = _build_document_entries(data)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)   # عرض A4 (بدون تغيير)
    section.page_height = Mm(203)  # طول مخصّص أقصر من A4 حسب طلبك (20.3 سم)
    section.left_margin = Pt(LEFT_MARGIN_PT)
    section.right_margin = Pt(LEFT_MARGIN_PT)
    section.top_margin = Pt(TOP_MARGIN_PT)
    section.bottom_margin = Pt(LEFT_MARGIN_PT)

    for text, size in entries:
        _mono_paragraph(doc, text, size)

    return doc


def generate_cd_document(data):
    """يبني المستند ويحفظه كملف Word نهائي، ويرجّع مساره."""
    doc = _build_docx(data)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    safe_passager = "".join(c for c in data["passager"] if c.isalnum() or c in " _-").strip() or "sans_nom"
    now = datetime.now()
    filename = f"CD_{safe_passager}_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(out_path)
    return out_path


# --- سجل مواضع الحقول القابلة للتعديل، فوق صورة النموذج مباشرة ---
# كل حقل: (رقم السطر بقائمة _build_document_entries، عمود بداية القيمة
# بالحرف، عرض الخانة بالحروف). حجم الخط يُستنتج من رقم السطر (نفس منطق
# _build_document_entries: 8-24 خط 11، الباقي خط 10).
FIELD_LAYOUT = {
    "no": (0, 54, 12),   # فراغين قبل الرقم أصلاً جزء من النص الأصلي "No  " — حده الأقصى الحقيقي 12 رقم
    "date": (3, 54, 13),
    # عمود "time" هنا (70) قيمة احتياطية غير مستخدمة فعلياً — cd_tab.py
    # يحسب مكانه الحقيقي حياً بالبكسل تبعاً لحرف "a" (راجع SplitDateEntry
    # .a_right_edge_px)، مُستخدم منها بس رقم السطر (3) لحساب الارتفاع y.
    "time": (3, 70, 6),
    "agence": (8, 15, 32),
    "guichet": (10, GUICHET_VALUE_COL, GUICHET_FIELD_WIDTH),
    "caisse": (11, 15, 32),
    "guichetier": (12, 15, 15),
    "passager": (15, NOM_REMETTANT_COL, 32),
    "passport_no": (16, PASSPORT_VALUE_COL, PASSPORT_FIELD_WIDTH),
    "date_delivrance": (16, DATE_DELIVRANCE_COL, 12),
    "eur": (18, EUR_COL, EUR_FIELD_WIDTH),
    "taux": (18, TAUX_COL, TAUX_FIELD_WIDTH),
    # Soit (dzd): نفس عمود وعرض حقل Montant en devise (EUR) بالضبط —
    # الاثنان مشتقّان من نفس الثوابت عمداً، بطلب صريح "نفس السعة
    # والخصائص"، حتى ما ينحرفوا عن بعض لو تغيّرت لاحقاً.
    "dzd": (19, EUR_COL, EUR_FIELD_WIDTH),
    "net_crediter": (24, 24, 20),
    "guichet_mirror": (10, GUICHET_NAME_COL, 32),
}

PAGE_WIDTH_PT = 210 * 2.8346456692913385
PAGE_HEIGHT_PT = 203 * 2.8346456692913385


def _line_font_size(index):
    return FONT_SIZE_MIDDLE if 8 <= index <= 24 else FONT_SIZE


def _line_height_for(index):
    return LINE_HEIGHT_MIDDLE_PT if _line_font_size(index) == FONT_SIZE_MIDDLE else LINE_HEIGHT_PT


def _y_top_pt(index):
    """أعلى نقطة (بالنقطة PDF) لسطر رقمه index، من أعلى الصفحة."""
    return TOP_MARGIN_PT + sum(_line_height_for(i) for i in range(index))


def field_layout_px(target_width_px):
    """
    يرجّع dict: اسم الحقل -> (x, y, width, height) بالبكسل، عند عرض صورة
    مقاسه target_width_px (نفس مقاس الصورة المعروضة بالواجهة).

    ملاحظة: عمود "time" هنا احتياطي فقط — cd_tab.py يحسب مكانه الحقيقي
    حياً بالبكسل من date_entry.a_right_edge_px() (راجع FIELD_LAYOUT).
    """
    scale = target_width_px / PAGE_WIDTH_PT
    result = {}
    for name, (index, col, width_chars) in FIELD_LAYOUT.items():
        size = _line_font_size(index)
        char_w = 0.6 * size
        x_pt = LEFT_MARGIN_PT + col * char_w
        y_pt = _y_top_pt(index)
        w_pt = width_chars * char_w
        h_pt = _line_height_for(index)
        result[name] = (x_pt * scale, y_pt * scale, w_pt * scale, h_pt * scale)
    return result


# --- ملفات المعاينة المؤقتة (تُستبدل في كل ضغطة "معاينة"، ما تتراكم) ---
PREVIEW_DOCX = os.path.join(OUTPUT_DIR, "_preview.docx")
PREVIEW_PDF = os.path.join(OUTPUT_DIR, "_preview.pdf")
PREVIEW_PNG = os.path.join(OUTPUT_DIR, "_preview.png")


def render_preview_image(data, target_width_px=750, out_png=None):
    """
    يبني المستند فعلياً، يحوّله PDF عن طريق Word نفسه (نفس ما يطبعه
    المستخدم بالضبط)، ثم يرسمه كصورة PNG حقيقية للمعاينة.
    يرجّع مسار ملف الـ PNG.
    """
    out_png = out_png or PREVIEW_PNG
    doc = _build_docx(data)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(PREVIEW_DOCX)

    import win32com.client as win32

    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    try:
        wdoc = word.Documents.Open(PREVIEW_DOCX)
        wdoc.SaveAs(PREVIEW_PDF, FileFormat=17)  # wdFormatPDF
        wdoc.Close(False)
    finally:
        word.Quit()

    import fitz

    pdf = fitz.open(PREVIEW_PDF)
    page = pdf[0]
    zoom = target_width_px / page.rect.width
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
    pix.save(out_png)
    pdf.close()

    return out_png


_EMPTY_DATA = {
    "no": "", "date": None, "time": "", "agence": "", "guichet": "", "caisse": "",
    "guichetier": "", "passager": "", "passport_no": "", "date_delivrance": None,
    "taux": None, "eur": None, "dzd": None,
    "_omit_a_literal": True,  # الخلفية الفاضية بس — راجع الشرح أعلاه
}


def get_blank_background(target_width_px=750, force=False):
    """
    يرسم صورة النموذج الفاضي (بس الأشياء الثابتة: Motif، PSP، Commission/
    Frais/Taxe = 0، الصندوق تحت) بنفس الدقة المطلوبة (target_width_px)
    ويخزّنها بملف خاص بهالدقة — حتى تكون الكتابة حادة دائماً بأي مستوى
    زوم (رسم أصلي بكل دقة، مو تكبير صورة صغيرة فتصير مبكسلة/مضببة).
    كل دقة تتولّد مرة وحدة بس (أول استخدام لها)، وتُعاد استخدامها بعدها.
    """
    path = os.path.join(OUTPUT_DIR, f"_blank_bg_{target_width_px}.png")
    if not force and os.path.exists(path):
        return path
    return render_preview_image(_EMPTY_DATA, target_width_px=target_width_px, out_png=path)
